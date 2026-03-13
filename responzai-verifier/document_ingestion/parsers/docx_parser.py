# document_ingestion/parsers/docx_parser.py

from docx import Document

async def parse_docx(file_path: str) -> dict:
    """
    Extrahiert Text aus Word-Dokumenten.

    Word-Dokumente haben eine klare Struktur:
    Überschriften, Absätze, Tabellen, Kopfzeilen, Fußzeilen.
    Wir behalten diese Struktur bei, weil sie für die
    Agenten wertvoll ist. Simon kann so erkennen, in welchem
    Kapitel eine Behauptung steht.
    """
    doc = Document(file_path)
    sections = []
    full_text = ""
    current_section = {"heading": "", "text": ""}

    for paragraph in doc.paragraphs:
        # Überschriften erkennen
        if paragraph.style.name.startswith("Heading"):
            # Vorherige Sektion speichern
            if current_section["text"].strip():
                sections.append(current_section.copy())
            current_section = {
                "heading": paragraph.text,
                "level": int(paragraph.style.name.replace("Heading ", "") or "1"),
                "text": ""
            }
        else:
            current_section["text"] += paragraph.text + "\n"

        full_text += paragraph.text + "\n"

    # Letzte Sektion speichern
    if current_section["text"].strip():
        sections.append(current_section)

    # Tabellen extrahieren
    for table in doc.tables:
        table_text = "\n"
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_text += " | ".join(cells) + "\n"
        full_text += table_text

    return {
        "text": full_text,
        "sections": sections,
        "pages": len(doc.sections)
    }
